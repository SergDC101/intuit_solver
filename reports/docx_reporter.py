import re

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from pathlib import Path

from data.env_var import ABSOLUT_PATH, GROUP_NAME, STUDENT_NAME


class DocxReporter:

    def __init__(self, course_name: str):
        self.course_name = course_name
        self.student_name = STUDENT_NAME
        self.group = GROUP_NAME

        self.screenshot_dir = Path(ABSOLUT_PATH + f"/{course_name}/скриншоты")
        # self.screenshot_dir.mkdir(parents=True, exist_ok=True)

        self.docs_dir = Path(ABSOLUT_PATH + f"/{course_name}/docx")
        self.docs_dir.mkdir(parents=True, exist_ok=True)

    def _natural_sort_key(self, filename):
        """
        Ключ для натуральной сортировки - извлекает числа из названия
        """
        # Ищем все числа в названии файла
        numbers = re.findall(r'\d+', filename.stem)
        if numbers:
            # Возвращаем первое найденное число для сортировки
            return int(numbers[0])
        else:
            # Если чисел нет, сортируем по названию
            return filename.stem

    def create_exam_document(self):
        # Создаем новый документ
        doc = Document()

        # Добавляем верхний колонтитул
        section = doc.sections[0]
        header = section.header
        header_paragraph = header.paragraphs[0]
        header_paragraph.text = f"{self.course_name} - ЭКЗАМЕН"
        header_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Добавляем нижний колонтитул
        footer = section.footer
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.text = f"{self.student_name}, {self.group}"
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


        # Получаем список скриншотов
        screenshot_files = []
        for file in Path(self.screenshot_dir).iterdir():
            if file.suffix.lower() in ['.png']:
                screenshot_files.append(file)

        # Сортируем файлы по имени
        screenshot_files.sort(key=self._natural_sort_key)


        # Добавляем скриншоты в документ
        for i, screenshot_path in enumerate(screenshot_files, 1):
            # Добавляем сам скриншот
            try:
                doc.add_picture(str(screenshot_path), width=Inches(6.0))
            except Exception as e:
                print(f"Ошибка при добавлении изображения {screenshot_path}: {e}")
                continue

            # Добавляем пустую строку между скриншотами
            doc.add_paragraph()

        # Сохраняем документ
        doc.save(str(self.docs_dir / "Экзамен.docx"))
        print(f"Документ успешно создан: {str(self.docs_dir / "Экзамен.docx")}")

